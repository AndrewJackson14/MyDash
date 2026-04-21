"""
Text extractors for press release attachments.

Two formats covered:
  - PDF: pdfplumber (handles most layouts, including columns)
  - DOCX: python-docx (handles Word documents)

Plain text and HTML email bodies are handled directly in bot.py
(no extraction needed for those).

Each extractor takes raw bytes and returns a single string. If
extraction fails for any reason, returns None and logs a warning —
the agent treats that release as text-only (uses just the email body
or the file's filename as a fallback).
"""
import io
import logging

logger = logging.getLogger(__name__)


def extract_pdf_text(pdf_bytes: bytes) -> str | None:
    """Extract text from a PDF byte stream. Returns None on failure."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed; cannot extract PDF text")
        return None

    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            chunks = []
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    chunks.append(txt)
            full = "\n\n".join(chunks).strip()
            if not full:
                logger.warning("PDF contained no extractable text (likely scanned image)")
                return None
            return full
    except Exception as e:
        logger.warning(f"PDF extraction failed: {e}")
        return None


def extract_docx_text(docx_bytes: bytes) -> str | None:
    """Extract text from a DOCX byte stream. Returns None on failure."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed; cannot extract DOCX text")
        return None

    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also pull cell text from any tables (press releases sometimes
        # use tables for letterhead headers).
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full = "\n\n".join(paragraphs).strip()
        if not full:
            return None
        return full
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
        return None


def extract_by_mime(content_bytes: bytes, mime_type: str, filename: str = "") -> str | None:
    """
    Dispatch to the right extractor based on MIME type or filename
    extension. Returns the extracted text, or None if no extractor
    handled the format.
    """
    mime = (mime_type or "").lower()
    name = (filename or "").lower()

    if mime == "application/pdf" or name.endswith(".pdf"):
        return extract_pdf_text(content_bytes)
    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or name.endswith((".docx", ".doc")):
        return extract_docx_text(content_bytes)
    if mime.startswith("text/"):
        try:
            return content_bytes.decode("utf-8", errors="ignore").strip() or None
        except Exception:
            return None

    logger.info(f"No extractor for mime={mime} filename={filename}")
    return None
